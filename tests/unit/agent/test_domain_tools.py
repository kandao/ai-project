"""
Unit tests for domain tools in agent/tools/

Tests: stocks (3.1), database (3.2), CSV analysis (3.3),
       chart generation (3.4), PDF/DOCX extraction (3.5).
"""

import json
import os
import sqlite3
import tempfile
from io import BytesIO
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest


# ── 3.1 Stock Price ──────────────────────────────────────────────────────

class TestStockPrice:

    def test_valid_ticker_returns_info(self):
        """3.1.1: Valid ticker → returns formatted price info string."""
        from tools.stocks import get_stock_price

        mock_ticker = MagicMock()
        mock_ticker.info = {
            "regularMarketPrice": 150.25,
            "shortName": "Apple Inc.",
            "regularMarketChange": 1.50,
            "regularMarketChangePercent": 1.01,
            "dayHigh": 151.00,
            "dayLow": 149.00,
            "volume": 50000000,
        }

        with patch("yfinance.Ticker", return_value=mock_ticker):
            result = get_stock_price("AAPL")

        assert "Apple Inc." in result
        assert "150.25" in result
        assert "AAPL" in result

    def test_invalid_ticker_no_price_returns_error(self):
        """3.1.2: Ticker with no price data → error message, no crash."""
        from tools.stocks import get_stock_price

        mock_ticker = MagicMock()
        mock_ticker.info = {}  # no price fields

        with patch("yfinance.Ticker", return_value=mock_ticker):
            result = get_stock_price("ZZZZZZZ")

        assert "Error" in result
        assert "ZZZZZZZ" in result

    def test_yfinance_exception_returns_error(self):
        """yfinance raises → graceful error string."""
        from tools.stocks import get_stock_price

        with patch("yfinance.Ticker", side_effect=Exception("network error")):
            result = get_stock_price("AAPL")

        assert "Error" in result

    def test_currentprice_fallback(self):
        """Uses currentPrice when regularMarketPrice is absent."""
        from tools.stocks import get_stock_price

        mock_ticker = MagicMock()
        mock_ticker.info = {
            "currentPrice": 200.00,
            "shortName": "Test Corp",
            "regularMarketChange": 0,
            "regularMarketChangePercent": 0,
            "dayHigh": 201.00,
            "dayLow": 199.00,
            "volume": 1000,
        }

        with patch("yfinance.Ticker", return_value=mock_ticker):
            result = get_stock_price("TST")

        assert "200.00" in result


# ── 3.2 Database Query ───────────────────────────────────────────────────

class TestDatabaseQuery:
    """DATABASE_URL env var is cleared so all tests use the SQLite fallback."""

    @pytest.fixture(autouse=True)
    def clear_db_url(self, monkeypatch):
        monkeypatch.delenv("DATABASE_URL", raising=False)

    def test_sqlite_select_returns_rows(self):
        """3.2.1: SELECT * FROM products → pipe-delimited table with column headers."""
        from tools.database import query_database

        result = query_database("SELECT * FROM products")
        # Verify pipe-delimited format with expected column names
        assert "|" in result
        assert "name" in result
        assert "price" in result

    def test_sqlite_no_results(self):
        """3.2.2: SELECT with impossible WHERE → 'No results.'"""
        from tools.database import query_database

        result = query_database("SELECT * FROM products WHERE price > 99999")
        assert result == "No results."

    def test_sqlite_invalid_sql(self):
        """3.2.3: Malformed SQL → first hits SELECT check or SQL Error."""
        from tools.database import query_database

        result = query_database("SELCET * FORM products")
        assert "Error" in result

    def test_only_select_allowed(self):
        """Non-SELECT statement → blocked."""
        from tools.database import query_database

        result = query_database("INSERT INTO products (name) VALUES ('hack')")
        assert "Error" in result
        assert "SELECT" in result

    def test_write_patterns_blocked(self):
        """SELECT containing DROP → blocked by write pattern check."""
        from tools.database import query_database

        result = query_database("SELECT * FROM products; DROP TABLE products")
        assert "Error" in result

    def test_dangerous_patterns_blocked(self):
        """UNION SELECT → blocked."""
        from tools.database import query_database

        result = query_database("SELECT * FROM products UNION SELECT * FROM employees")
        assert "Error" in result

    def test_query_too_long_blocked(self):
        """Query > 2000 chars → blocked."""
        from tools.database import query_database

        result = query_database("SELECT " + "a" * 2000)
        assert "Error" in result
        assert "too long" in result

    def test_result_truncation_at_100_rows(self):
        """3.2.5: >100 rows → first 100 + '(N total rows)' note."""
        from tools.database import _format_rows

        cols = ["id", "val"]
        rows = [(i, i * 2) for i in range(150)]
        result = _format_rows(cols, rows, 150)
        assert "150 total rows" in result

    def test_postgres_mode_with_mock(self):
        """3.2.4: db_url provided → _query_postgres called."""
        from tools.database import query_database

        with patch("tools.database._query_postgres", return_value="1") as mock_pg:
            result = query_database("SELECT 1", db_url="postgresql://user:pass@localhost/db")

        mock_pg.assert_called_once()
        assert result == "1"

    def test_auto_create_sample_db(self, tmp_path):
        """3.2.6: Delete sample.db, call query → DB recreated."""
        from tools import database as db_mod

        orig_path = db_mod.DB_PATH
        db_mod.DB_PATH = tmp_path / "sample.db"
        try:
            assert not db_mod.DB_PATH.exists()
            result = db_mod.query_database("SELECT COUNT(*) FROM products")
            assert db_mod.DB_PATH.exists()
            assert "Error" not in result
        finally:
            db_mod.DB_PATH = orig_path


# ── 3.3 CSV Analysis ─────────────────────────────────────────────────────

class TestCSVAnalysis:

    @pytest.fixture
    def csv_file(self, tmp_path):
        f = tmp_path / "data.csv"
        f.write_text("name,price,qty\nWidget,9.99,100\nGadget,19.99,50\nDoohickey,4.99,200\n")
        return str(f)

    def test_basic_csv_analysis(self, csv_file):
        """3.3.1: Valid CSV → shape, columns, dtypes, head, describe."""
        from tools import _analyze_csv

        result = _analyze_csv(csv_file)
        assert "Shape:" in result
        assert "Columns:" in result
        assert "price" in result
        assert "Widget" in result

    def test_csv_with_query(self, csv_file):
        """3.3.2: query='price.mean()' → query result appended."""
        from tools import _analyze_csv

        result = _analyze_csv(csv_file, query="price.mean()")
        assert "Query result" in result

    def test_csv_file_not_found(self):
        """3.3.3: Non-existent path → 'Error: ...'"""
        from tools import _analyze_csv

        result = _analyze_csv("/nonexistent/path/data.csv")
        assert result.startswith("Error:")

    def test_csv_invalid_query(self, csv_file):
        """3.3.4: Invalid pandas expression → 'Query error: ...'"""
        from tools import _analyze_csv

        result = _analyze_csv(csv_file, query="this_is_not_valid!!!")
        assert "Query error" in result or "Error" in result


# ── 3.4 Chart Generation ─────────────────────────────────────────────────

class TestChartGeneration:

    @pytest.fixture
    def chart_data(self):
        return json.dumps({"Q1": 100, "Q2": 150, "Q3": 120, "Q4": 200})

    def test_bar_chart_saved(self, tmp_path, chart_data):
        """3.4.1: bar chart → file saved at output_path."""
        from tools import _generate_chart

        out = str(tmp_path / "bar.png")
        result = _generate_chart(chart_data, chart_type="bar", output_path=out)
        assert "Chart saved to" in result
        assert Path(out).exists()

    def test_line_chart_saved(self, tmp_path, chart_data):
        """3.4.2: line chart → file saved."""
        from tools import _generate_chart

        out = str(tmp_path / "line.png")
        result = _generate_chart(chart_data, chart_type="line", output_path=out)
        assert "Chart saved to" in result
        assert Path(out).exists()

    def test_pie_chart_saved(self, tmp_path, chart_data):
        """3.4.3: pie chart → file saved."""
        from tools import _generate_chart

        out = str(tmp_path / "pie.png")
        result = _generate_chart(chart_data, chart_type="pie", output_path=out)
        assert "Chart saved to" in result
        assert Path(out).exists()

    def test_invalid_json_returns_error(self, tmp_path):
        """3.4.4: Malformed JSON → 'Error: ...'"""
        from tools import _generate_chart

        out = str(tmp_path / "bad.png")
        result = _generate_chart("not valid json {{{", output_path=out)
        assert "Error" in result

    def test_unknown_chart_type_falls_back_to_bar(self, tmp_path, chart_data):
        """Unknown chart_type → falls back to bar, file still saved."""
        from tools import _generate_chart

        out = str(tmp_path / "fallback.png")
        result = _generate_chart(chart_data, chart_type="radar", output_path=out)
        assert "Chart saved to" in result


# ── 3.5 PDF / DOCX Extraction ────────────────────────────────────────────

class TestPDFExtraction:

    def test_extract_valid_pdf(self, tmp_path):
        """3.5.1: PDF with text → text content returned."""
        from tools.pdf_extractor import extract_pdf

        # Create a minimal valid PDF using reportlab if available, else skip
        pytest.importorskip("reportlab")
        from reportlab.pdfgen import canvas

        pdf_path = str(tmp_path / "test.pdf")
        c = canvas.Canvas(pdf_path)
        c.drawString(100, 750, "Hello from PDF")
        c.save()

        result = extract_pdf(pdf_path)
        assert "Hello from PDF" in result

    def test_extract_invalid_file_returns_error(self, tmp_path):
        """3.5.3: Random bytes as PDF → 'Error: ...'"""
        from tools.pdf_extractor import extract_pdf

        bad_pdf = tmp_path / "bad.pdf"
        bad_pdf.write_bytes(b"\x00\x01\x02\x03 not a pdf")
        result = extract_pdf(str(bad_pdf))
        assert result.startswith("Error:")

    def test_extract_pdf_not_found(self):
        """Non-existent PDF → 'Error: ...'"""
        from tools.pdf_extractor import extract_pdf

        result = extract_pdf("/nonexistent/file.pdf")
        assert result.startswith("Error:")


class TestDOCXExtraction:

    def test_extract_valid_docx(self, tmp_path):
        """3.5.2: DOCX with paragraphs → paragraph text returned."""
        from tools.doc_extractor import extract_doc
        from docx import Document

        docx_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Hello from DOCX")
        doc.add_paragraph("Second paragraph")
        doc.save(str(docx_path))

        result = extract_doc(str(docx_path))
        assert "Hello from DOCX" in result
        assert "Second paragraph" in result

    def test_extract_empty_docx(self, tmp_path):
        """DOCX with no paragraphs → 'No text found in document.'"""
        from tools.doc_extractor import extract_doc
        from docx import Document

        docx_path = tmp_path / "empty.docx"
        doc = Document()
        doc.save(str(docx_path))

        result = extract_doc(str(docx_path))
        assert "No text found" in result

    def test_extract_invalid_file_returns_error(self, tmp_path):
        """3.5.3: Random bytes as DOCX → 'Error: ...'"""
        from tools.doc_extractor import extract_doc

        bad_docx = tmp_path / "bad.docx"
        bad_docx.write_bytes(b"not a docx file at all")
        result = extract_doc(str(bad_docx))
        assert result.startswith("Error:")

    def test_extract_doc_not_found(self):
        """Non-existent DOCX → 'Error: ...'"""
        from tools.doc_extractor import extract_doc

        result = extract_doc("/nonexistent/file.docx")
        assert result.startswith("Error:")
